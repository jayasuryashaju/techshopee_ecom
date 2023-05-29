from io import BytesIO
from docx import Document

from django.core.paginator import Paginator
from django.views import View
from xhtml2pdf import pisa
from django.db.models import Count, Q
from django.template.loader import get_template
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone
from datetime import datetime
import datetime
from .forms import UserForm, CouponForm
from account.models import User
from products.models import Category
from products.forms import CategoryForm
from products.models import Product, Banner
from products.forms import ProductForm, BannerForm
from account.forms import RegistrationForm
from cart.models import Order, Coupon, CategoryOffer, ProductOffer, OrderItem
from cart.forms import CategoryOfferForm, ProductOfferForm


# Create your views here.


def admin_panel(request):
    username = request.session.get('username')
    if username:
        user = User.objects.all()
        orders = Order.objects.all()
        products = Product.objects.all()
        categories = Category.objects.all()
        number_of_products = products.count()
        number_of_categories = categories.count()
        number_of_orders = orders.count()
        total_revenue = 0
        for order in orders:
            total_revenue += order.total_price

        today = datetime.datetime.today()
        month_start = today.replace(day=1)
        next_month_start = (month_start + datetime.timedelta(days=32)).replace(day=1)
        orders_this_month = Order.objects.filter(created_at__gte=month_start, created_at__lt=next_month_start)
        total_revenue_this_month = 0
        for order in orders_this_month:
            total_revenue_this_month += order.total_price

        today = datetime.datetime.now()
        dates = Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(
            order_items=Count('id')).order_by('created_at__day')

        returns = Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(
            retunrs=Count('id', filter=Q(status='returned'))).order_by('created_at__day')

        sales = Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(
            sales=Count('id', filter=Q(status='completed'))).order_by('created_at__day')

        most_moving_product_count = []
        most_moving_product = []
        products = Product.objects.all()
        for i in products:
            most_moving_product.append(i)
            most_moving_product_count.append(OrderItem.objects.filter(product=i).count())


        placed_count = Order.objects.filter(status='pending').count()
        shipped_count = Order.objects.filter(status='shipped').count()
        delivered_count = Order.objects.filter(status='completed').count()
        return_count = Order.objects.filter(status='returned').count()
        cancelled_count = Order.objects.filter(status='cancelled').count()
        context = {
            'orders': orders,
            'user': user,
            'number_of_products': number_of_products,
            'number_of_categories': number_of_categories,
            'number_of_orders': number_of_orders,
            'total_revenue': total_revenue,
            'total_revenue_this_month': total_revenue_this_month,
            'sales': sales,
            'returns': returns,
            'dates': dates,
            'most_moving_product': most_moving_product,
            'most_moving_product_count': most_moving_product_count,
            'status_count': [
                placed_count,
                shipped_count,
                delivered_count,
                return_count,
                cancelled_count,
            ]
        }
        return render(request, "admin/dashboard.html", context)
    return redirect('ecom_admin:admin_login')


def user_details(request):
    username = request.session.get('username')
    if username:
        user = User.objects.all()
        return render(request, "admin/userview.html", {'user': user})
    return redirect('ecom_admin:admin_login')


def admin_login(request):
    username = request.session.get('username')
    if username:
        return redirect('ecom_admin:admin_panel')
    global user
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(username=username, password=password)
        print(username, password)
        print(user)
        # if request.user.is_authenticated:
        #     login(request, user)
        if user.is_authenticated:
            login(request, user)
            print(user.is_authenticated)
            print(user.is_superuser)
            if user.is_superuser:
                request.session['username'] = username
                return redirect('ecom_admin:admin_panel')
            else:
                logout(request)
                request.session.flush()
                messages.error(request, 'user is not an admin')
                return redirect('ecom_admin:admin_login')
        else:
            messages.error(request, 'username or password is incorrect')
    return render(request, 'admin_login.html')


@login_required(login_url='admin_login')
def edit_user(request, pk):
    user = User.objects.get(id=pk)
    uform = UserForm(instance=user)

    if request.method == 'POST':
        uform = UserForm(request.POST, instance=user)
        print(uform.errors)
        if not uform.errors:
            if uform.is_valid():
                uform.save()
                return redirect('ecom_admin:admin_panel')
        else:
            messages.error(request, 'Some filed are already in use')
            context = {
                'uform': uform
            }
            return render(request, 'admin/useredit.html', context)

    context = {
        'uform': uform
    }
    return render(request, 'admin/useredit.html', context)


@login_required(login_url='admin_login')
def block_user(request, pk):
    user = User.objects.get(id=pk)
    print(pk)
    if request.method == 'GET':
        print(pk)
        if user.is_active:
            user.is_active = False
            user.save()
        else:
            user.is_active = True
            user.save()

        return redirect('ecom_admin:admin_panel')
    return redirect('ecom_admin:admin_panel')


@login_required(login_url='admin_login')
def delete_user(request, username):
    u = User.objects.get(username=username)
    u.delete()
    messages.success(request, 'User deleted successfully')
    return redirect('ecom_admin:admin_panel')


def admin_logout(request):
    if request.user.is_authenticated:
        logout(request)
        request.session.flush()
        messages.error(request, 'User is not an Admin')
        return redirect('ecom_admin:admin_login')
    #     ab = False
    #     ab = user.is_authenticated
    #     if ab == True:
    #         dict_val = {'dictval': User.objects.all()}
    #         return render(request, 'admin/userview.html', dict_val)
    # else:
    #     logout(request)
    #     request.session.flush()
    #     messages.error(request, 'User is not an Admin')
    #     return redirect('ecom_admin:admin_login')


@login_required(login_url='admin_login')
def add_user(request):
    pass
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            first_name = form.cleaned_data['first_name']
            last_name = form.cleaned_data['last_name']
            phone_number = form.cleaned_data['phone_number']
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            username = email.split("@")[0]
            user = User.objects.create_user(
                first_name=first_name, last_name=last_name, email=email, username=username, password=password,
                phone_number=phone_number)
            user.save()
            return redirect('account:user_signin')
    else:
        form = RegistrationForm()
    context = {
        'form': form,
    }
    return render(request, 'signup.html', context)


@login_required(login_url='admin_login')
def categories_view(request):
    categories = Category.objects.all()
    context = {
        'categories': categories
    }
    return render(request, 'admin/categories.html', context)


@login_required(login_url='admin_login')
def add_category(request):
    if request.method == 'POST':
        form = CategoryForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('ecom_admin:categories_view')
    else:
        form = CategoryForm()
    context = {
        'form': form
    }
    return render(request, 'admin/add_category.html', context)


@login_required(login_url='admin_login')
def edit_category(request, pk):
    category = get_object_or_404(Category, pk=pk)
    if request.method == 'POST':
        form = CategoryForm(request.POST, request.FILES, instance=category)
        if form.is_valid():
            form.save()
            return redirect('ecom_admin:categories_view')
    else:
        form = CategoryForm(instance=category)
    return render(request, 'admin/edit_category.html', {'form': form, 'category': category})


@login_required(login_url='admin_login')
def delete_category(request, pk):
    category = get_object_or_404(Category, pk=pk)
    if request.method == 'POST':
        category.delete()
        return redirect('ecom_admin:categories_view')
    return render(request, 'admin/delete_category.html', {'category': category})


# def product_list(request):
#     products = Product.objects.all()
#     return render(request, 'admin/product_list.html', {'products': products})
@login_required(login_url='admin_login')
def product_list(request):
    categories = Category.objects.all()
    products = Product.objects.all()
    category_id = request.GET.get('category')
    if category_id:
        products = products.filter(category_id=category_id)

    if request.method == 'POST':
        product_id = request.POST.get('product_id')
        product = Product.objects.get(id=product_id)
        product.delete()
        return redirect('product_list')

    context = {
        'categories': categories,
        'products': products
    }
    return render(request, 'admin/product_list.html', context)


@login_required(login_url='admin_login')
def add_product(request, category_id=None):
    category = None
    if category_id:
        category = Category.objects.get(id=category_id)
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            product = form.save(commit=False)
            product.category = form.cleaned_data['category']
            product.save()
            return redirect('ecom_admin:product_list')
    else:
        form = ProductForm(selected_category=category)
    return render(request, 'admin/product_add.html', {'form': form, 'category': category})


@login_required(login_url='admin_login')
def edit_product(request, pk):
    # Retrieve the product object with the given pk
    product = get_object_or_404(Product, pk=pk)
    categories = Category.objects.all()

    if request.method == 'POST':
        # Create a form instance with the submitted data and the product instance
        form = ProductForm(request.POST, request.FILES, instance=product)
        if form.is_valid():
            # Save the updated form data and redirect to the product detail page
            form.save()
            return redirect('ecom_admin:product_list')
        else:
            print(form.errors)
    else:
        # Create a form instance with the product instance
        form = ProductForm(instance=product)

    return render(request, 'admin/edit_product.html', {'product': product, 'categories': categories, 'form': form})


@login_required(login_url='admin_login')
def delete_product(request, pk):
    # Retrieve the product object with the given pk
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        # Delete the product object and redirect to the product list page
        product.delete()
        return redirect('ecom_admin:product_list')
    return render(request, 'admin/delete_product.html', {'product': product})


@login_required(login_url='admin_login')
def product_details(request, pk):
    product = get_object_or_404(Product, pk=pk)
    return render(request, 'admin/product_details.html', {'product': product})


@login_required(login_url='admin_login')
def add_coupon(request):
    if request.method == 'POST':
        form = CouponForm(request.POST)
        if form.is_valid():
            print("hello the for is valid")
            coupon = form.save(commit=False)
            coupon.created = timezone.now()
            coupon.save()
            messages.success(request, 'Coupon added successfully')
            return redirect('ecom_admin:coupon_list')
    else:
        form = CouponForm()
    return render(request, 'admin/create_coupon.html', {'form': form})


def coupon_list(request):
    coupons = Coupon.objects.filter(active=True)
    return render(request, 'admin/coupon_list.html', {'coupons': coupons})


def edit_coupon(request, pk):
    coupon = get_object_or_404(Coupon, id=pk)
    form = CouponForm(request.POST or None, instance=coupon)
    if form.is_valid():
        form.save()
        return redirect('ecom_admin:coupon_list')
    return render(request, 'admin/create_coupon.html', {'form': form})


def delete_coupon(request, pk):
    coupon = get_object_or_404(Coupon, id=pk)
    coupon.delete()
    return redirect('ecom_admin:coupon_list')
    # return render(request, 'admin/coupon_list.html', {'coupon': coupon})


def offers(request):
    category_offers = CategoryOffer.objects.all()
    product_offers = ProductOffer.objects.all()
    context = {
        'category_offers': category_offers,
        'product_offers': product_offers,
    }
    return render(request, 'admin/offers.html', context)


def add_category_offer(request):
    form = CategoryOfferForm()
    if request.method == 'POST':
        form = CategoryOfferForm(request.POST)
        discount = form.data['discount']
        if int(discount) <= 70:
            if form.is_valid():
                form.save()
                return redirect('ecom_admin:offers')
            else:

                messages.error(request, 'Already exist!!')
                return redirect('ecom_admin:add_category_offer')
        else:
            messages.error(request, 'Percentage should be less than or equal to 70')
            return redirect('add_category_offer')
    return render(request, 'admin/add_offer.html', {'form': form, })


def edit_category_offer(request, category_id):
    print('hello i am edit offer')
    category_offer = CategoryOffer.objects.get(pk=category_id)
    form = None
    if request.method == 'POST':
        form = CategoryOfferForm(request.POST, instance=category_offer)

        if form.is_valid():
            form.save()
            return redirect("ecom_admin:offers")
    else:
        form = CategoryOfferForm(instance=category_offer)

    context = {'form': form, 'category_id': category_id}
    return render(request, 'admin/edit_offer.html', context)


def apply_category_offer(request, category_offer_id):
    category_offer = CategoryOffer.objects.get(pk=category_offer_id)
    user_id = request.user
    if not user_id:
        return redirect("ecom_admin:login")
    if category_offer.is_active:
        return redirect("ecom_admin:offers")
    print(category_offer.category)
    products = Product.objects.filter(category=category_offer.category)
    all_offers = CategoryOffer.objects.filter(category=category_offer.category)
    for offer in all_offers:
        if offer.is_active:
            deactivate_offer(request, offer.id)
            print('hello i am offer id', offer.id)

    for product in products:
        if product.selling_price != 0:
            if float(product.selling_price) < float(product.price) - (
                    float(product.price) * (category_offer.discount / 100)):
                continue
        else:
            product.selling_price = float(product.price) - (float(product.price) * (category_offer.discount / 100))
            product.save()
    category_offer.is_active = True
    category_offer.save()
    return redirect("ecom_admin:offers")


def deactivate_offer(request, category_offer_id):
    category_offer = CategoryOffer.objects.get(pk=category_offer_id)
    user_id = request.user
    if not user_id:
        return redirect("ecom_admin:login")
    if not category_offer.is_active:
        return redirect("ecom_admin:offers")
    category_offer.is_active = False
    category_offer.save()
    products = Product.objects.filter(category=category_offer.category)
    for product in products:
        # product_offer = ProductOffer.objects.filter(product=product)
        # if len(product_offer) != 0:
        #     print(product_offer.discount)
        # # if product_offer.product.name == product.name:
        # #     product.selling_price = float(product.price) - (float(product.price) * (product_offer.discount / 100))
        # #     product.save()
        # # else:
        product.selling_price = 0
        product.save()

    return redirect("ecom_admin:offers")


def delete_category_offer(request, id):
    cat_offer = get_object_or_404(CategoryOffer, id=id)
    category_offer_id = cat_offer.category.id
    deactivate_offer(request, category_offer_id)
    cat_offer.delete()
    messages.error(request, 'Deleted')
    return redirect('ecom_admin:offers')


def add_product_offer(request):
    form = ProductOfferForm()
    if request.method == 'POST':
        form = ProductOfferForm(request.POST)
        discount = form.data['discount']
        if int(discount) <= 70:
            if form.is_valid():
                form.save()
                return redirect('ecom_admin:offers')
            else:
                messages.error(request, 'Already exist!!')
                return redirect('ecom_admin:add_product_offer')
        else:
            messages.error(request, 'Percentage should be less than or equal to 70')
            return redirect('add_product_offer')
    return render(request, 'admin/add_offer.html', {'form': form, })


def apply_product_offer(request, product_offer_id):
    product_offer = ProductOffer.objects.get(pk=product_offer_id)
    product_category = product_offer.product.category
    category_offers = CategoryOffer.objects.filter(category=product_category)

    user_id = request.user
    if not user_id:
        return redirect("ecom_admin:login")
    if product_offer.is_active:
        return redirect("ecom_admin:offers")
    print(product_offer.product.category)
    product = Product.objects.get(id=product_offer.product.id)
    for offer in category_offers:
        if offer.is_active:
            print('hai')
            if offer.discount > product_offer.discount:
                return redirect("ecom_admin:offers")
            else:
                product.selling_price = float(product.price) - (float(product.price) * (product_offer.discount / 100))
                product.save()
                product_offer.is_active = True
                product_offer.save()
    product.selling_price = float(product.price) - (float(product.price) * (product_offer.discount / 100))
    product.save()
    product_offer.is_active = True
    product_offer.save()
    return redirect("ecom_admin:offers")


def delete_product_offer(request, id):
    product_offer = get_object_or_404(ProductOffer, id=id)
    product_offer_id = product_offer.id
    deactivate_product_offer(request, product_offer_id)
    product_offer.delete()
    messages.error(request, 'Deleted')
    return redirect('ecom_admin:offers')


def edit_product_offer(request, product_id):
    print('hello i am edit offer')
    product_offer = ProductOffer.objects.get(pk=product_id)
    form = None
    if request.method == 'POST':
        form = ProductOfferForm(request.POST, instance=product_offer)

        if form.is_valid():
            form.save()
            return redirect("ecom_admin:offers")
    else:
        form = ProductOfferForm(instance=product_offer)

    context = {'form': form, 'product_id': product_id}
    return render(request, 'admin/edit_offer.html', context)


def deactivate_product_offer(request, product_offer_id):
    product_offer = ProductOffer.objects.get(pk=product_offer_id)
    user_id = request.user
    if not user_id:
        return redirect("ecom_admin:login")
    if not product_offer.is_active:
        return redirect("ecom_admin:offers")
    product = Product.objects.get(id=product_offer.product.id)
    category_offers = CategoryOffer.objects.filter(category=product.category)
    for offer in category_offers:
        if offer.is_active:
            product.selling_price = float(product.price) - (float(product.price) * (offer.discount / 100))
            product.save()
            product_offer.is_active = False
            product_offer.save()
            return redirect("ecom_admin:offers")
    product.selling_price = 0
    product.save()
    product_offer.is_active = False
    product_offer.save()
    return redirect("ecom_admin:offers")


@login_required(login_url='admin_login')
def order_details(request):
    orders = Order.objects.all()

    context = {
        'orders': orders,
    }
    return render(request, 'admin/order_details.html', context)


@login_required(login_url='admin_login')
def order_details_product(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    order_items = OrderItem.objects.filter(order_id=order_id)

    context = {
        'order': order,
        'order_items': order_items,

    }

    return render(request, 'admin/order_details_product.html', context)


def change_order_status(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    if request.method == 'POST':
        new_status = request.POST.get('status')
        print('hei i am new status', new_status)
        order.status = new_status
        order.save()
        # return HttpResponseRedirect(reverse('ecom_admin:order_details_product', args=(order_id,)))
        return redirect('ecom_admin:order_details_product', order_id=order_id)
    return render(request, 'admin/order_details_product.html', {'order': order})


# sales_report


def render_to_pdf(template_src, context_dict={}):
    template = get_template(template_src)
    html = template.render(context_dict)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)
    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return None


def sales_report(request):
    orders = Order.objects.all()
    new_order_list = []

    for i in orders:
        order_items = OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id': i.id,
                'ordered_date': i.created_at,
                'user': i.user,
                'price': j.price,
                'method': i.payment_method,
                'status': i.status,
            }

            new_order_list.append(item)
    paginator = Paginator(new_order_list, 3)
    page = request.GET.get('page')
    paged_orders_list = paginator.get_page(page)
    return render(request, 'admin/sales-report.html', {
        'order': paged_orders_list
    })


def by_date(request):
    if request.GET.get('from'):
        sales_date_from = datetime.datetime.strptime(request.GET.get('from'), "%Y-%m-%d")
        sales_date_to = datetime.datetime.strptime(request.GET.get('to'), "%Y-%m-%d")

        sales_date_to += datetime.timedelta(days=1)
        orders = Order.objects.filter(created_at__range=[sales_date_from, sales_date_to])

        new_order_list = []

        for i in orders:
            order_items = OrderItem.objects.filter(order_id=i.id)
            for j in order_items:
                item = {
                    'id': i.id,
                    'ordered_date': i.created_at,
                    'user': i.user,
                    'price': j.price,
                    'method': i.payment_method,
                    'status': i.status
                }
                new_order_list.append(item)
    else:
        messages.error(request, 'Select fields before submitting...!!!')
        return redirect('ecom_admin:sales_report')

    return render(request, 'admin/sales-report.html', {
        'order': new_order_list,
    })


class generates_sales_report(View):
    def get(self, request, *args, **kwargs):
        try:
            orders = Order.objects.all()
            new_order_list = []

            for i in orders:
                order_items = OrderItem.objects.filter(order_id=i.id)
                for j in order_items:
                    item = {
                        'id': i.id,
                        'ordered_date': i.created_at,
                        'user': i.user.full_name,
                        'price': j.price,
                        'method': i.payment_method,
                        'status': i.status
                    }
                    new_order_list.append(item)
        except:
            return HttpResponse("505 not found")

        data = {
            'order': new_order_list
        }

        pdf = render_to_pdf('admin/sales-report-pdf.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


def by_month(request):
    month = request.GET.get('month')
    orders = Order.objects.filter(created_at__month=month)
    new_order_list = []

    for i in orders:
        order_items = OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id': i.id,
                'ordered_date': i.created_at,
                'user': i.user,
                'price': j.price,
                'method': i.payment_mode,
                'status': i.status,
            }
            new_order_list.append(item)

    return render(request, 'admin/sales-report.html', {
        'order': new_order_list,
    })


def by_year(request):
    year = request.GET.get('year')
    orders = Order.objects.filter(created_at__year=year)
    new_order_list = []

    for i in orders:
        order_items = OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id': i.id,
                'ordered_date': i.created_at,
                'user': i.user,
                'price': j.price,
                'method': i.payment_method,
                'status': i.status
            }
            new_order_list.append(item)

    return render(request, 'admin/sales-report.html', {
        'order': new_order_list,
    })


def download_docx(request):
    document = Document()
    document.add_heading('Sales Report', 0)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocuments.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=sales-report.docx'
    document.save(response)

    return response


def banner_list(request):
    banners = Banner.objects.all()
    return render(request, 'admin/banners.html', {'banners': banners})


def banner_add(request):
    products = Product.objects.all()  # Retrieve all products
    if request.method == 'POST':
        form = BannerForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('ecom_admin:banner_list')
    else:
        form = BannerForm()
    return render(request, 'admin/banner_add.html', {'form': form, 'products': products})


def banner_edit(request, banner_id):
    banner = get_object_or_404(Banner, id=banner_id)
    if request.method == 'POST':
        form = BannerForm(request.POST, request.FILES, instance=banner)
        if form.is_valid():
            form.save()
            return redirect('ecom_admin:banner_list')
    else:
        form = BannerForm(instance=banner)
    return render(request, 'admin/banner_add.html', {'form': form, 'banner': banner})


def banner_delete(request, banner_id):
    banner = get_object_or_404(Banner, id=banner_id)
    if request.method == 'POST':
        banner.delete()
        return redirect('ecom_admin:banner_list')
